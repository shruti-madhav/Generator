# Imports
import os
import time
import openpyxl
import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Loading the Excel file
excel_file = 'Excel_File.xlsx'
wb = openpyxl.load_workbook(excel_file, data_only=True)
sheet = wb.active

# Loading the Word file
template_file = 'PO_TEMPLATE.docx'
doc_template = Document(template_file)

# Directory to save generated Word documents with excel data
output_dir = 'generated_pos'
os.makedirs(output_dir, exist_ok=True)

# Getting column indices for excel data 
# Make sure these are same as excel file
# sr_col = 1  
purchase_indent_date_col = 2       
purchase_indent_no_col = 3   
purchase_order_date_col = 4
delivery_date_col = 5
purchase_order_number_col = 6
vendor_name_col = 7
vendor_address_col = 8
hsn_sac_col = 9
currency_code_col = 10
exchange_rate_col = 11
reference_no_col = 12
account_col = 13
item_price_col = 14
item_desc_col = 15
quantity_ordered_col = 16
uom_col = 17
item_total_col = 18
igst_col = 19
cgst_col =20
sgst_col =21
contact_no_col=22
email_col = 23

# Dictionary to store items per PO
po_dict = {}

# Going through the rows in the Excel sheet
for row in sheet.iter_rows(min_row=2, max_row=20, values_only=True):  # For skipping the header row  
    po_number = row[purchase_order_number_col - 1]
    po_date = row[purchase_order_date_col - 1].date().strftime('%d/%m/%Y')
    vendor_name = row[vendor_name_col - 1]
    vendor_address = row[vendor_address_col - 1]
    item_price = row[item_price_col - 1]
    item_total = row[item_total_col - 1]
    item_desc = row[item_desc_col - 1]
    hsn_sac = row[hsn_sac_col - 1]
    reference_no = row[reference_no_col-1]
    currency_code = row[currency_code_col-1]
    quantity_ordered = row[quantity_ordered_col - 1]
    uom = row[uom_col - 1]
    delivery_date = row[delivery_date_col - 1].date().strftime('%d/%m/%Y')
    igst = row[igst_col-1]
    cgst =row[cgst_col-1]
    sgst = row[sgst_col-1]
    contact_no = row[contact_no_col-1]
    email = row[email_col-1]
     
    if delivery_date == "":
        delivery_date = " NA "
    else:
        delivery_date = row[delivery_date_col - 1].date().strftime('%d/%m/%Y')
        # delivery_date = row[delivery_date_col - 1]
    
    # Store items for each PO
    if po_number not in po_dict:
        po_dict[po_number] = []
    
    po_dict[po_number].append(
        {
            'po_number': po_number, 
            'vendor_name': vendor_name,
            'po_date': po_date,
            'item_price': item_price,
            'item_total': item_total,
            'item_desc': item_desc,
            'hsn_sac': hsn_sac,
            'reference_no': reference_no,
            'currency_code':currency_code,
            'quantity_ordered': quantity_ordered,
            'uom': uom,
            'delivery_date': delivery_date,
            'vendor_address': vendor_address,
            'igst':igst,
            'sgst':sgst,
            'cgst':cgst,
            'contact_no':contact_no,
            'email':email
        }
    )

# Function to add borders to a table cell
def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"val": "single", "sz": "12", "color": "#000000", "space": "0"},
        bottom={"val": "single", "sz": "12", "color": "#000000", "space": "0"},
        left={"val": "single", "sz": "12", "color": "#000000", "space": "0"},
        right={"val": "single", "sz": "12", "color": "#000000", "space": "0"},
    )
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    for border_name, border_attr in kwargs.items():
        element = OxmlElement(f"w:{border_name}")
        for attr_name, attr_value in border_attr.items():
            element.set(qn(f"w:{attr_name}"), str(attr_value))
        tcPr.append(element)
    for paragraph in cell.paragraphs:
        # paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Center the text
        for run in paragraph.runs:
            run.font.size = Pt(8)



def number_to_words_inr(number):
    ones = ['', 'one', 'two', 'three', 'four', 'five', 'six', 'seven', 'eight', 'nine']
    teens = ['ten', 'eleven', 'twelve', 'thirteen', 'fourteen', 'fifteen', 'sixteen', 'seventeen', 'eighteen', 'nineteen']
    tens = ['', '', 'twenty', 'thirty', 'forty', 'fifty', 'sixty', 'seventy', 'eighty', 'ninety']
    thousands = ['', 'thousand', 'lakh', 'crore']

    def convert_number_to_words(n):
        if n == 0:
            return 'zero'

        words = []

        if n >= 10000000:
            words.append(convert_number_to_words(n // 10000000))
            words.append('crore')
            n %= 10000000

        if n >= 100000:
            words.append(convert_number_to_words(n // 100000))
            words.append('lakh')
            n %= 100000

        if n >= 1000:
            words.append(convert_number_to_words(n // 1000))
            words.append('thousand')
            n %= 1000

        if n >= 100:
            words.append(convert_number_to_words(n // 100))
            words.append('hundred')
            n %= 100

        if n >= 20:
            words.append(tens[n // 10])
            n %= 10

        if n >= 10:
            words.append(teens[n - 10])
            n = 0

        if n > 0:
            words.append(ones[n])

        return ' '.join(words).strip()

    return convert_number_to_words(int(number)) + ' rupees'



# Generating Word documents for each PO
for po_number, items in po_dict.items():
    PO_NO = po_number.replace('/', '').replace('-', '')
    total_amount = 0
    total_sgst = 0
    total_cgst = 0
    total_igst = 0
    total_tax = 0
    total_po = 0
    for elements in items:
        total_amount += elements['item_total']
        total_cgst += elements['cgst']
        total_igst += elements['igst']
        total_sgst += elements['sgst']
        if 'Telangana' not in elements['vendor_address']: 
            total_po = total_amount + total_igst
            total_tax = total_igst
        else :
            total_po = total_amount + total_sgst+total_cgst
            total_tax = total_sgst+total_cgst
    
    # Create a new document based on the template
    doc = Document(template_file)
    time.sleep(0.3)
    # Iterate through all tables in the document
    for table in doc.tables:
        # Iterate through each row in the table
        for row in table.rows:
            # Iterate through each cell in the row
            for cell in row.cells:
                # Check and replace placeholders in the cell text
                if '<<PO_NUMBER>>' in cell.text:
                    cell.text = cell.text.replace('<<PO_NUMBER>>', str(items[0]['po_number'])) 
                    cell.bold = True
                if '<<PO_DATE>>' in cell.text:
                    cell.text = cell.text.replace('<<PO_DATE>>', str(items[0]['po_date']))
                    cell.bold = True
                if '<<VENDOR_NAME>>' in cell.text:
                    cell.text = cell.text.replace('<<VENDOR_NAME>>', str(items[0]['vendor_name']))
                if '<<VENDOR_ADDRESS>>' in cell.text:
                    cell.text = cell.text.replace('<<VENDOR_ADDRESS>>', str(items[0]['vendor_address']))
                if '<<ITEM_PRICE>>' in cell.text:
                    cell.text = cell.text.replace('<<ITEM_PRICE>>', str(items[0]['item_price']))
                if '<<ITEM_TOTAL>>' in cell.text:
                    cell.text = cell.text.replace('<<ITEM_TOTAL>>', str(items[0]['item_total']))
                if '<<TOTAL_AMOUNT>>' in cell.text:
                    cell.text = cell.text.replace('<<TOTAL_AMOUNT>>',"₹ "+str(total_amount))
                
                total_amount_words=number_to_words_inr(total_amount).capitalize()
                if '<<TOTAL_AMOUNT_WORDS>>' in cell.text:
                    cell.text = cell.text.replace('<<TOTAL_AMOUNT_WORDS>>',total_amount_words)
                if '<<CURRENCY_CODE>>' in cell.text:
                    cell.text = cell.text.replace('<<CURRENCY_CODE>>',str(items[0]['currency_code']))
                if '<<REFERENCE_NO>>' in cell.text:
                    cell.text = cell.text.replace('<<REFERENCE_NO>>',str(items[0]['reference_no']))
                else:
                    cell.text = cell.text.replace('<<REFERENCE_NO>>','NA')
                if 'Telangana' not in items[0]['vendor_address']:   
                    if '<<TOTAL_IGST>>' in cell.text:
                        cell.text = cell.text.replace('<<TOTAL_IGST>>', "₹ "+str(items[0]['igst']))
                    if '<<TOTAL_CGST>>' in cell.text:
                        cell.text = cell.text.replace('<<TOTAL_CGST>>','')
                    if '<<TOTAL_SGST>>' in cell.text:
                        cell.text = cell.text.replace('<<TOTAL_SGST>>','')
                else :
                    if '<<TOTAL_IGST>>' in cell.text:
                        cell.text = cell.text.replace('<<TOTAL_IGST>>', '')
                    if '<<TOTAL_CGST>>' in cell.text:
                        cell.text = cell.text.replace('<<TOTAL_CGST>>',"₹ "+str(items[0]['cgst']))
                    if '<<TOTAL_SGST>>' in cell.text:
                        cell.text = cell.text.replace('<<TOTAL_SGST>>',"₹ "+str(items[0]['sgst']))
                if'<<TOTAL_PO>>' in cell.text:
                    cell.text = cell.text.replace('<<TOTAL_PO>>',"₹ "+str(total_po))
                if '<<TOTAL_TAX>>' in cell.text:
                    cell.text = cell.text.replace('<<TOTAL_TAX>>',"₹ "+str(total_tax) )
                
                if '<<CONTACT_NO>>' in cell.text:
                    contact_no = items[0]['contact_no']
                    cell.text = cell.text.replace('<<CONTACT_NO>>', str(contact_no) if contact_no is not None else " ")

                if '<<EMAIL>>' in cell.text:
                    email = items[0]['email']
                    cell.text = cell.text.replace('<<EMAIL>>', str(email) if email is not None else " ")

                for paragraph in cell.paragraphs:
                    # paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Center the text
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        # For adding new rows in the table
        srno = 0
        for item in items:
            srno+=1
            row_cells = table.add_row().cells
           
            if len(row_cells) < 13:    
                continue  

            row_cells[0].text = str(srno)
            row_cells[1].text = str(item['item_desc'])
            row_cells[2].text = str(item['hsn_sac'])
            row_cells[3].text = str(item['quantity_ordered'])
            row_cells[4].text = str(item['uom'])
            if item['currency_code'] == "INR":
                row_cells[5].text = "₹"+str(item['item_price'])
            if item['currency_code'] == "USD": 
                row_cells[5].text = "$"+str(item['item_price'])
                
            row_cells[6].text = str(item['delivery_date'])
            row_cells[7].text = str(" NA ")  
            row_cells[8].text = str(" NA ")  
            row_cells[9].text = str(item['item_total'])
            
            if 'Telangana' not in item['vendor_address']:
                row_cells[10].text = str(item['igst'])
                row_cells[11].text = '-'
                row_cells[12].text = '-'

            else:
                row_cells[10].text = '-'
                row_cells[11].text = str(item['sgst'])
                row_cells[12].text = str(item['cgst'])

            # Apply borders to each cell in the newly added row
            for cell in row_cells:

                set_cell_border(cell,
                    top={"val": "single", "sz": "6", "color": "000000", "space": "0"},
                    bottom={"val": "single", "sz": "6", "color": "000000", "space": "0"},
                    left={"val": "single", "sz": "6", "color": "000000", "space": "0"},
                    right={"val": "single", "sz": "6", "color": "000000", "space": "0"},
                )
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the text
                    for run in paragraph.runs:
                        run.font.size = Pt(8)  


    # Saving the generated Word docs
    output_path = os.path.join(output_dir, f'{PO_NO}.docx')
    print("This is the last path: " + output_path)
    doc.save(output_path)

print("PO documents generated successfully!")
